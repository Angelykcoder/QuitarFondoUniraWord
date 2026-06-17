import copy
import io
import logging
import os
import time
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.table import Table, _Cell
from PIL import Image

# ==========================================
# CONFIGURACIÓN DE RUTAS
# ==========================================
TEMPLATE_DOCX = r"C:\Users\LENOVO\Desktop\GENERADOR IMPRESIONES\FormatoImprimir.docx" #AQUI SE COLOCA LA PLANTILLA CON TAMAÑOS
IMAGES_FOLDER = r"C:\Users\LENOVO\Desktop\Resultado"  #AQUI SE ESPECIFICA DE DONDE TOMO LAS FOTOS
OUTPUT_DOCX = r"C:\Users\LENOVO\Desktop\OUT\STAGE\Fotos_Finales.docx"   #AQUI LA CARPETA DONDE SE PONDRAN LA FOTOS

# Configuración del Sistema de Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - [%(levelname)s] - %(message)s",
    handlers=[
        logging.FileHandler("automatizacion_word.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logging.getLogger("docx").setLevel(logging.WARNING)
logging.getLogger("PIL").setLevel(logging.WARNING)


def asegurar_plantilla_valida(template_path: Path, ancho_cm: float, alto_cm: float) -> None:
    """
    Fuerza la creación de una plantilla optimizada. 
    Se eliminó la validación simple para REESCRIBIR el molde con los nuevos márgenes máximos.
    """
    logging.info("Configurando/Optimizando plantilla con márgenes máximos de impresión...")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # Reducimos los márgenes a 1 cm en todo el documento para ganar espacio máximo
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        
    # Nueva estructura: 6 columnas de ancho (entran dos grupos de 3) y 5 filas de alto
    COLUMNAS = 6
    FILAS = 5
    
    tabla = doc.add_table(rows=FILAS, cols=COLUMNAS, style='Table Grid')
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in tabla.rows:
        row.height = Cm(alto_cm)
        for cell in row.cells:
            cell.width = Cm(ancho_cm)
            cell.text = ""
            
    doc.save(str(template_path))
    logging.info(f"Plantilla optimizada generada en: {template_path}")


def resize_fit_with_padding(image_path: Path, target_width_cm: float, target_height_cm: float, dpi: int = 300) -> io.BytesIO:
    """
    Redimensiona la foto de forma proporcional para que quepa exacta sin recortarse.
    """
    target_width_px = int((target_width_cm / 2.54) * dpi)
    target_height_px = int((target_height_cm / 2.54) * dpi)

    with Image.open(image_path) as img:
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
            
        lienzo_blanco = Image.new("RGB", (target_width_px, target_height_px), (255, 255, 255))
        img.thumbnail((target_width_px, target_height_px), Image.Resampling.LANCZOS)
        
        pos_x = (target_width_px - img.width) // 2
        pos_y = (target_height_px - img.height) // 2
        
        lienzo_blanco.paste(img, (pos_x, pos_y))
        
        img_byte_arr = io.BytesIO()
        lienzo_blanco.save(img_byte_arr, format='JPEG', quality=90)
        img_byte_arr.seek(0)
        
        return img_byte_arr


def obtener_celdas_tabla(table: Table) -> List[_Cell]:
    """Saca todas las celdas de una tabla en una lista lineal."""
    return [cell for row in table.rows for cell in row.cells]


def main() -> None:
    start_time = time.time()
    
    # Tus dimensiones estrictas deseadas
    ANCHO_CM = 2.6
    ALTO_CM = 3.9
    
    template_path = Path(TEMPLATE_DOCX)
    images_dir = Path(IMAGES_FOLDER)
    output_path = Path(OUTPUT_DOCX)
    
    # Forzar la actualización del archivo base con el nuevo diseño extendido
    asegurar_plantilla_valida(template_path, ANCHO_CM, ALTO_CM)

    if not images_dir.exists() or not images_dir.is_dir():
        logging.error(f"La carpeta de imágenes no es válida o no existe: {images_dir}")
        print(f"Error: Carpeta de imágenes no encontrada en {images_dir}")
        return

    valid_extensions = {'.jpg', '.jpeg', '.png', '.bmp'}
    image_files = sorted([
        f for f in images_dir.iterdir() 
        if f.is_file() and f.suffix.lower() in valid_extensions
    ])
    
    total_imagenes = len(image_files)
    if total_imagenes == 0:
        logging.warning("No se encontraron imágenes aptas para procesar.")
        print(f"Imágenes encontradas: 0 en la ruta {images_dir}")
        return

    print(f"Imágenes encontradas: {total_imagenes}")
    logging.info(f"Iniciando procesamiento optimizado de {total_imagenes} imágenes.")

    try:
        doc = Document(str(template_path))
        template_table = doc.tables[0]
        
        # Cargar las 30 celdas iniciales disponibles en la nueva página max-espacio
        celdas_disponibles = obtener_celdas_tabla(template_table)
        
        puntero_celda = 0
        paginas_generadas = 1

        for idx, img_path in enumerate(image_files, 1):
            print(f"[{idx}/{total_imagenes}] Procesando {img_path.name}")
            
            try:
                img_stream = resize_fit_with_padding(img_path, ANCHO_CM, ALTO_CM, dpi=300)
                
                # Insertar en bloques de 3 celdas continuas
                for _ in range(3):
                    if puntero_celda >= len(celdas_disponibles):
                        doc.add_page_break()
                        
                        cloned_tbl_xml = copy.deepcopy(template_table._tbl)
                        doc.element.body.append(cloned_tbl_xml)
                        
                        nueva_tabla = Table(cloned_tbl_xml, doc)
                        celdas_disponibles.extend(obtener_celdas_tabla(nueva_tabla))
                        paginas_generadas += 1
                    
                    celda_actual = celdas_disponibles[puntero_celda]
                    
                    celda_actual.text = ""
                    parrafo = celda_actual.paragraphs[0]
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = parrafo.add_run()
                    img_stream.seek(0)
                    run.add_picture(img_stream, width=Cm(ANCHO_CM), height=Cm(ALTO_CM))
                    
                    puntero_celda += 1
                
                img_stream.close()
                
            except Exception as e:
                logging.error(f"Fallo al procesar la imagen {img_path.name}. Detalle: {e}")
                print(f"   [!] Error omitido en imagen: {img_path.name}")
                continue

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        elapsed_time = time.time() - start_time
        tiempo_formateado = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
        
        print("...")
        print(f"Páginas generadas: {paginas_generadas}")
        print(f"Tiempo transcurrido: {tiempo_formateado}")
        print("Documento guardado y optimizado correctamente.")
        logging.info(f"Proceso exitoso. Guardado en: {output_path}")

    except Exception as e:
        logging.critical(f"Error catastrófico en la ejecución: {e}", exc_info=True)
        print(f"\nOcurrió un error grave durante la ejecución: {e}")


if __name__ == "__main__":
    main()